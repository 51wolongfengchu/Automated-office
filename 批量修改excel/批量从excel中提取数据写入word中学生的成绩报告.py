import pandas as pd
from docx import Document
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

#读取所需要的表格
df = pd.read_excel('期末考试成绩.xlsx')
df = df.fillna(0)
#`axis=1` 表示沿着行的方向进行求和，即对每一行的对应值进行求和。
#.loc` 属性进行行列索引，选取了从列名 "yu wen" 到列名 "sehng wu"（包括这两列）的所有列。`:` 表示选取所有行
df["总分"] = df.loc[:,"语文":"生物"].sum(axis=1)
#rank` 方法对 DataFrame 的每一行进行排序，method="min"` 表示使用最小排名方法，ascending=False` 表示按照降序进行排序
#astype` 方法将排序后的结果转换为整数类型，即将排名从浮点数转换为整数。
df["期末排名"] = df["总分"].rank(method="min",ascending=False).astype("int")
#int将数据类型转换为整数类型
df["期中排名"] = df["期中排名"].astype("int")
#进步的名次=期末-其中
df["进步名次"] = -(df["期末排名"]-df["期中排名"])

# 打开指定的文档
doc = Document('模板.docx')

# 调用normal_run函数来修改文档中具体位置的段落和运行对象的样式
def normal_run(a,b,text):
    # 获取指定位置的段落和运行对象
    run = doc.paragraphs[a].runs[b]
    # 设置运行对象的文本内容为传入的text参数
    run.text = str(text)
    # 将运行对象的字体样式设置为加粗
    run.font.bold = True
    # 将运行对象的字体颜色设置为RGB值为(2, 30, 170)
    run.font.color.rgb = RGBColor(2,30,170)


def table_run(c, text):
    # 获取要修改的单元格对象，`c` 参数表示要修改的单元格列的索引
    cell = doc.tables[0].rows[1].cells[c]
    # 获取单元格中的段落对象
    paragraph = cell.paragraphs[0]
    # 将段落的对齐方式设置为居中对齐
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 获取段落中的运行对象
    run = paragraph.runs[0]
    # 设置运行对象的文本内容为传入的text参数
    run.text = str(text)
    # 将运行对象的字体样式设置为加粗
    run.font.bold = True
    # 将运行对象的字体大小设置为12磅
    run.font.size = Pt(12)

#使用 `iterrows()` 方法循环迭代 DataFrame `df` 的每一行，同时获取每一行的索引和数据
for index,rows in df.iterrows():
    print(rows["姓名"])
    #调用`normal_run` 函数，将第0行第0列的单元格内容设置为当前行的 "1" 列的值。
    normal_run(0, 0, rows["班级"])
    # 设置第0行第2列的单元格内容为当前行的 "2" 列的值
    normal_run(0, 2, rows["姓名"])
    normal_run(1, 2, rows["总分"])
    normal_run(1, 7, rows["进步名次"])
    normal_run(2, 4, rows["期中排名"])
    normal_run(2, 9, rows["期末排名"])
    table_run(1, rows["语文"])
    table_run(2, rows["数学"])
    table_run(3, rows["英语"])
    table_run(4, rows["政治"])
    table_run(5, rows["历史"])
    table_run(6, rows["地理"])
    table_run(7, rows["物理"])
    table_run(8, rows["化学"])
    # 设置第9行的单元格内容为当前行的 "90" 列的值
    table_run(9, rows["生物"])
    # 将修改后的Word文档保存到以当前行的 "姓名" 列的值命名的文件中

    # 指定文件夹路径
    folder_path = './教育'
    # 检查文件夹是否存在.exists()方法表示是否存在
    if not os.path.exists(folder_path):
        # 若文件夹不存在，则创建文件夹
        os.mkdir(folder_path)
        print("文件夹创建成功")
    else:
        print("文件夹已存在")
    doc.save(f'./教育/{rows["姓名"]}的成长足迹.docx')