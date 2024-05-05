import os
import docx
import pandas as pd
from docx.document import Document
from docx.oxml.text import run
from docx.text import paragraph

#获取文件夹
file_list = os.listdir("./合同文件")
#读取并筛选出符合条件的合同文件
new_file_list = []
for index,file in enumerate(file_list):
    if file.startswith('服务合同-'):
        new_file_list.append(file_list[index])
#一定注意路径，格式
full_path_list = ["./合同文件/" + i for i in new_file_list]

#可以改导入excel中的要写入的内容
data = pd.DataFrame(columns=("我是傻逼", "合作名称", "委托公司", "签订日期", "合同金额", "合作名称"))

for full_path in full_path_list:
    #读取每个合同的文件
    doc = docx.Document(full_path)
    result = []
    #读取文档中第一段第二个文字块的颜色，只是文字块的颜色是一样的(可以修改文字块字体类型，字体大小等在书上200页）
    rgb = doc.paragraphs[0].runs[1].font.size

    for paragraph in doc.paragraphs:
        tmp = []
        for run in paragraph.runs:
        #筛选出颜色一致的文字块并存储到列表中
            #if用来检查，（str用来把文字块转换为字符串）是否和rgb变量相等，相等的话就添加到tmp的列表中
            if str(run.font.size) == str(rgb):
                tmp.append(run.text)
        if tmp:
            #将 "tmp" 列表中的文字块连接成一个字符串，其中的空格表示字体间的间隔。然后，将该字符串添加到结果列表 "result" 中。
            result.append("".join(tmp))
        #将得到的数据转换为数据框，可以修改columns的数据，导入excel文档中的内容也就会被修改
    df = pd.DataFrame([result], columns=("我是傻逼", "合作名称", "委托公司", "签订日期", "合同金额", "合作名称"))
    #将得到的小数据框合并成一个大的数据框
    data = data._append(df)
#将所有的数据写入本地的excel文件中
data.to_excel("合同文件信息_导出.xlsx", index=None)

print('have finished')

#注意文字块要和正文有所区别，文字块改改颜色，字体等，和正文进行区分