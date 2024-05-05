import pandas as pd
from docx import Document
from docx.shared import RGBColor

#python-docx---读取docx文件
doc = Document('通知单模板.docx')
#pandas----读取xlsx表
df = pd.read_excel('物业费明细表.xlsx')
#遍历读取表中的每一行数据
for index,rows in df.iterrows():
    print(index,rows[0],rows[1],rows[2])
def style(run):
    #字体设置为是否加粗
    run.font.bold = True
    #字体设置为是否加下划线
    run.font.underline = True
    #字体设置填充颜色为蓝色
    run.font.color.rgb = RGBColor(45,105,150)
for index,rows in df.iterrows():
    #第二段第一个需要填充的文字快
    run1 = doc.paragraphs[2].runs[1]
    run1.text = rows[0]

    #第三段第4个需要填充的文字块
    run2 = doc.paragraphs[3].runs[4]
    run2.text = rows[1]
    style(run2)

    #第三段第24个需要填充的文字快
    run3 = doc.paragraphs[3].runs[24]
    run3.text = str(rows[2])
    style(run3)

    #写入文件名字
    doc.save(f'{rows[1]} - 文件的名字.docx')

print('have finished')