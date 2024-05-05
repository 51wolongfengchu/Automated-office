from docx import Document

#文件的名字
doc = Document("5_3_5.docx")

def replace_word(doc, old, new):
    #遍历每一个自然段
    for p in doc.paragraphs:
        #使用run保证不改变原文的样式。否则就会改变原文的所有样式
        for run in p.runs:
            #内容里面的旧文字和新文字
            run.text = run.text.replace(old, new)
    #遍历表格（是否有表格）
    for table in doc.tables:
        #每个表格中的行
        for row in table.rows:
            #每一个单元格里面的内容
            for cell in row.cells:
                #单元格里面的文字进行替换方法
                cell.text = cell.text.replace(old, new)

#所需要被替换的文字和替换的文字
replace_word(doc, "Python", "888")
#写入文件的名字
doc.save("5_3_5_替换后.docx")

print('have finished')